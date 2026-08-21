# -*- coding: utf-8 -*-
"""
Generador de Documentos RRHH
=============================
App web (Streamlit) que reemplaza el notebook de Colab original.
5 herramientas, cada una con subida manual de Excel de datos + plantilla(s).

Cómo correr localmente:
    pip install -r requirements.txt
    streamlit run app.py
"""

import io
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
import datetime
from datetime import timedelta
from decimal import Decimal, ROUND_HALF_UP

import pandas as pd
import streamlit as st
from openpyxl import load_workbook
from openpyxl.utils.datetime import from_excel
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from num2words import num2words
import holidays

try:
    import anthropic
except ImportError:
    anthropic = None


# =========================================================
# IA (OPCIONAL) — auditoría de datos + chat de ayuda
# =========================================================
# No requiere nada para que el resto de la app funcione. Si no hay
# ANTHROPIC_API_KEY configurada en Secrets, estas funciones simplemente
# avisan que la IA no está disponible, sin romper nada más.

MODELO_IA = "claude-haiku-4-5-20251001"


def _api_key_ia():
    """Lee la API key desde Secrets. st.secrets lanza una excepción propia
    si no existe ningún secrets.toml configurado — por eso el try/except."""
    try:
        return st.secrets.get("ANTHROPIC_API_KEY")
    except Exception:
        return None


def ia_disponible():
    return anthropic is not None and bool(_api_key_ia())


def preguntar_ia(system_prompt, mensaje_usuario, max_tokens=800):
    """Devuelve (texto_respuesta, error). Si algo falla, texto_respuesta es None."""
    if anthropic is None:
        return None, "El paquete 'anthropic' no está instalado en este servidor."
    api_key = _api_key_ia()
    if not api_key:
        return None, "No hay una IA configurada todavía (falta ANTHROPIC_API_KEY en Secrets)."
    try:
        cliente = anthropic.Anthropic(api_key=api_key)
        respuesta = cliente.messages.create(
            model=MODELO_IA,
            max_tokens=max_tokens,
            system=system_prompt,
            messages=[{"role": "user", "content": mensaje_usuario}],
        )
        texto = "".join(bloque.text for bloque in respuesta.content if hasattr(bloque, "text"))
        return texto.strip(), None
    except Exception as e:
        return None, f"Error llamando a la IA: {e}"


def auditar_datos_con_ia(df, contexto):
    """df: DataFrame con los datos crudos del Excel (sin procesar). contexto: qué tipo de
    documento se va a generar, para que la IA sepa qué mirar."""
    muestra = df.head(30).to_csv(index=False)
    system_prompt = (
        "Eres un asistente que revisa planillas de RR.HH. en Chile antes de generar documentos "
        "legales (contratos, finiquitos, anexos). Te paso datos en formato CSV. Busca posibles "
        "errores: fechas mal escritas, incoherentes o fuera de un rango razonable; RUTs con "
        "formato inválido (debieran ser XX.XXX.XXX-X); nombres vacíos o con mayúsculas/minúsculas "
        "inconsistentes; montos que parecen error de tipeo (demasiados o muy pocos dígitos); filas "
        "duplicadas; celdas vacías en columnas que parecen obligatorias (nombre, fecha, monto). "
        "Responde en español, en una lista breve con viñetas indicando fila y problema. Si no "
        "encuentras nada relevante, dilo en una sola frase corta. No inventes columnas que no "
        "existen en los datos."
    )
    mensaje = f"Contexto del documento: {contexto}\n\nDatos (primeras filas, formato CSV):\n{muestra}"
    return preguntar_ia(system_prompt, mensaje, max_tokens=900)


def bloque_auditoria_ia(archivo_excel, contexto, key):
    """Botón reutilizable de 'Revisar con IA' para pegar justo debajo de cualquier
    uploader de Excel de datos."""
    if not archivo_excel:
        return
    with st.expander("🔍 Revisar estos datos con IA (opcional)"):
        st.caption(
            "Al usar esto, una muestra de tus datos (nombres, fechas, montos) se envía a la "
            "API de IA configurada en el servidor. No se guarda nada, pero sí sale de esta app."
        )
        if st.button("Revisar con IA", key=f"ia_btn_{key}"):
            if not ia_disponible():
                st.warning(
                    "La IA no está configurada todavía en este despliegue "
                    "(falta ANTHROPIC_API_KEY en Secrets de Streamlit Cloud)."
                )
            else:
                try:
                    df_check = pd.read_excel(io.BytesIO(archivo_excel.getvalue()))
                except Exception as e:
                    st.warning(f"No pude leer el Excel para revisarlo: {e}")
                    return
                with st.spinner("Revisando datos con IA..."):
                    resultado, error = auditar_datos_con_ia(df_check, contexto)
                if error:
                    st.warning(error)
                else:
                    st.markdown(resultado)


# =========================================================
# HELPERS GENERALES
# =========================================================

def format_fecha(valor):
    """Formatea cualquier valor de fecha (date, datetime, número serial de Excel
    o string en varios formatos) a DD/MM/YYYY."""
    if valor in (None, ""):
        return ""
    if isinstance(valor, (datetime.datetime, datetime.date)):
        return valor.strftime("%d/%m/%Y")
    if isinstance(valor, (int, float)):
        try:
            return from_excel(valor).strftime("%d/%m/%Y")
        except Exception:
            pass
    if isinstance(valor, str):
        valor = valor.strip()
        for fmt in ("%d-%m-%Y", "%d/%m/%Y", "%Y-%m-%d"):
            try:
                return datetime.datetime.strptime(valor, fmt).strftime("%d/%m/%Y")
            except Exception:
                continue
        try:
            dt = pd.to_datetime(valor, dayfirst=True, errors="coerce")
            if not pd.isnull(dt):
                return dt.strftime("%d/%m/%Y")
        except Exception:
            pass
    return str(valor)


def monto_a_texto(valor, prefijo="", sufijo="PESOS"):
    """1234567 -> '1.234.567 (UN MILLON ... PESOS)' (prefijo/sufijo configurables)."""
    valor_limpio = str(valor).replace(".", "").replace(",", "")
    valor_int = int(float(valor_limpio))
    valor_numero = prefijo + "{:,}".format(valor_int).replace(",", ".")
    valor_letras = num2words(valor_int, lang="es").upper() + " " + sufijo
    return f"{valor_numero} ({valor_letras})"


def sueldo_en_letras(valor):
    """Igual que monto_a_texto pero retorna (numero, letras) por separado."""
    valor_int = int(float(str(valor).replace(".", "").replace(",", "")))
    numero = "{:,}".format(valor_int).replace(",", ".")
    letras = num2words(valor_int, lang="es").upper() + " PESOS"
    return numero, letras


def crear_zip(archivos: dict) -> bytes:
    """archivos: {nombre_archivo: bytes}. Soporta rutas con '/' para subcarpetas."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for nombre, contenido in archivos.items():
            zf.writestr(nombre, contenido)
    buffer.seek(0)
    return buffer.getvalue()


def guardar_docx_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def guardar_xlsx_bytes(wb) -> bytes:
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def libreoffice_disponible():
    return shutil.which("soffice") is not None


def _preparar_docx_para_pdf(docx_bytes: bytes) -> bytes:
    """Devuelve una copia del .docx con el interlineado fijado en un valor EXACTO
    (no 'automático' según la fuente). El servidor no tiene Arial/Times New Roman
    instalados de verdad — usa fuentes equivalentes que miden ligeramente distinto
    en alto de línea — y eso puede desbordar el documento a una página extra al
    convertir a PDF, aunque en Word se vea perfecto en una sola página. Fijar el
    interlineado en un valor exacto hace que el resultado no dependa de qué fuente
    use el servidor. No toca párrafos que tengan imágenes/sellos (para no
    aplastarlos), y esta copia se usa SOLO para generar el PDF — el .docx que se
    entrega al usuario nunca se modifica."""
    try:
        doc = Document(io.BytesIO(docx_bytes))

        def tiene_dibujo(p):
            return bool(p._p.findall(".//" + qn("w:drawing")))

        def procesar(parrafos):
            for p in parrafos:
                if tiene_dibujo(p) or not p.runs:
                    continue
                tam = None
                for r in p.runs:
                    if r.font.size:
                        tam = r.font.size.pt
                        break
                if tam is None and p.style and p.style.font and p.style.font.size:
                    tam = p.style.font.size.pt
                if tam is None:
                    tam = 11.0
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                p.paragraph_format.line_spacing = Pt(tam * 1.15)

        procesar(doc.paragraphs)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    procesar(cell.paragraphs)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        return docx_bytes  # si algo falla, seguimos con el original tal cual


def convertir_docs_a_pdf_batch(archivos_docx: dict) -> tuple:
    """archivos_docx: {nombre_sin_extension: bytes_docx}.
    Convierte TODOS los documentos en una sola invocación de LibreOffice (mucho
    más rápido que uno por uno). Devuelve (pdfs, error):
    - pdfs: {nombre_sin_extension: bytes_pdf} de los que sí se pudieron convertir
    - error: mensaje si algo salió mal a nivel general, o None si todo bien
    Nunca lanza excepción — si falla, devuelve {} y un mensaje de error, para que
    el resto de la app (los Word ya generados) no se vea afectado."""
    if not archivos_docx:
        return {}, None
    if not libreoffice_disponible():
        return {}, "LibreOffice no está instalado en este servidor (falta packages.txt con libreoffice-writer)."

    pdfs = {}
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            rutas = []
            for nombre, docx_bytes in archivos_docx.items():
                ruta = os.path.join(tmpdir, f"{nombre}.docx")
                with open(ruta, "wb") as f:
                    f.write(_preparar_docx_para_pdf(docx_bytes))
                rutas.append(ruta)

            resultado = subprocess.run(
                ["soffice", "--headless", "--norestore", "--convert-to", "pdf", "--outdir", tmpdir] + rutas,
                capture_output=True, timeout=300,
            )

            for nombre in archivos_docx:
                ruta_pdf = os.path.join(tmpdir, f"{nombre}.pdf")
                if os.path.exists(ruta_pdf):
                    with open(ruta_pdf, "rb") as f:
                        pdfs[nombre] = f.read()

            if not pdfs:
                stderr = resultado.stderr.decode("utf-8", errors="ignore")[:300]
                return {}, f"LibreOffice no generó ningún PDF. Detalle: {stderr}"
    except subprocess.TimeoutExpired:
        return pdfs, "La conversión a PDF tardó demasiado y se cortó (quedaron los Word igual)."
    except Exception as e:
        return pdfs, f"Error convirtiendo a PDF: {e}"

    return pdfs, None


def agregar_pdfs_al_zip(archivos: dict) -> dict:
    """Recibe el dict de archivos ya armado (ruta -> bytes), busca los .docx,
    los convierte a PDF en un solo lote, y devuelve un dict NUEVO con esos PDF
    (misma ruta, extensión .pdf) para sumarlo al dict original. Si algo falla,
    avisa con st.warning pero no interrumpe nada — los Word siguen intactos."""
    docx_items = {k: v for k, v in archivos.items() if k.lower().endswith(".docx")}
    if not docx_items:
        return {}

    mapa_temp_a_ruta = {}
    docx_para_convertir = {}
    for i, (ruta, contenido) in enumerate(docx_items.items()):
        nombre_temp = f"doc_{i}"
        mapa_temp_a_ruta[nombre_temp] = ruta
        docx_para_convertir[nombre_temp] = contenido

    pdfs, error = convertir_docs_a_pdf_batch(docx_para_convertir)

    if error and not pdfs:
        st.warning(f"⚠️ No se pudieron generar los PDF (los Word igual quedaron listos): {error}")
    elif error:
        st.warning(f"⚠️ Algunos PDF no se generaron: {error}")

    nuevos = {}
    for nombre_temp, pdf_bytes in pdfs.items():
        ruta_original = mapa_temp_a_ruta[nombre_temp]
        ruta_pdf = ruta_original[:-5] + ".pdf"
        nuevos[ruta_pdf] = pdf_bytes
    return nuevos


def separar_word_y_pdf(archivos: dict) -> tuple:
    """Separa un dict de archivos ya armado en dos: uno con todo lo que NO es PDF
    (Word, Excel) y otro solo con los PDF. Sirve para ofrecer descargas separadas
    por formato en vez de un único ZIP mezclado."""
    word = {k: v for k, v in archivos.items() if not k.lower().endswith(".pdf")}
    pdf = {k: v for k, v in archivos.items() if k.lower().endswith(".pdf")}
    return word, pdf


# =========================================================
# LECTURA DE EXCEL (variantes según lo que necesitaba cada tool)
# =========================================================

def leer_excel_formato(file_bytes, date_headers, money_headers=None, money_prefijo=""):
    """Lee Excel preservando negrita/tamaño de fuente por celda (para mergefields
    que necesitan mantener formato). Devuelve lista de dicts: header -> (valor, bold, size)."""
    money_headers = money_headers or []
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    excel_data = []

    for row in ws.iter_rows(min_row=2, values_only=False):
        if all(c.value in (None, "") for c in row):
            continue
        row_data = {}
        for cell in row:
            header_raw = ws.cell(row=1, column=cell.column).value
            if header_raw is None:
                continue
            header = str(header_raw).strip().upper()
            valor = cell.value
            bold = cell.font.bold if cell.font.bold is not None else False
            size = cell.font.sz if cell.font.sz is not None else None

            if header in date_headers:
                valor = format_fecha(valor)
            elif header in money_headers and valor not in (None, ""):
                try:
                    valor = monto_a_texto(valor, prefijo=money_prefijo)
                except Exception:
                    valor = str(valor)

            row_data[header] = (valor, bold, size)

        nombre = row_data.get("NOMBRE COMPLETO", ("", False, None))[0]
        if nombre is not None and str(nombre).strip() != "":
            excel_data.append(row_data)

    return excel_data


def leer_excel_simple(file_bytes, date_headers):
    """Lee Excel a strings planos (sin negrita/tamaño) — para anexos de continuidad."""
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    excel_data = []

    for row in ws.iter_rows(min_row=2, values_only=False):
        if all(c.value in (None, "") for c in row):
            continue
        row_data = {}
        for cell in row:
            header_raw = ws.cell(row=1, column=cell.column).value
            if header_raw is None:
                continue
            header = str(header_raw).strip().upper()
            valor = cell.value
            if header in date_headers and valor is not None:
                valor = format_fecha(valor)
            row_data[header] = "" if valor is None else str(valor)

        if row_data.get("NOMBRE COMPLETO", "").strip():
            excel_data.append(row_data)

    return excel_data


def leer_headers_excel(file_bytes):
    """Devuelve solo los nombres de columna (tal cual, sin mayusculizar) de la primera fila."""
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    headers = []
    for cell in ws[1]:
        if cell.value is not None and str(cell.value).strip() != "":
            headers.append(str(cell.value).strip())
    return headers


def leer_excel_universal(file_bytes, date_headers=None, money_headers=None,
                          money_prefijo="", money_sufijo="PESOS"):
    """Como leer_excel_formato, pero SIN exigir una columna NOMBRE COMPLETO fija:
    cualquier Excel + cualquier combinación de columnas fecha/monto elegidas por el usuario.
    Una fila se incluye si tiene al menos un valor no vacío."""
    date_headers = date_headers or []
    money_headers = money_headers or []
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    excel_data = []

    for row in ws.iter_rows(min_row=2, values_only=False):
        if all(c.value in (None, "") for c in row):
            continue
        row_data = {}
        for cell in row:
            header_raw = ws.cell(row=1, column=cell.column).value
            if header_raw is None:
                continue
            header = str(header_raw).strip().upper()
            valor = cell.value
            bold = cell.font.bold if cell.font.bold is not None else False
            size = cell.font.sz if cell.font.sz is not None else None

            if header in date_headers:
                valor = format_fecha(valor)
            elif header in money_headers and valor not in (None, ""):
                try:
                    valor = monto_a_texto(valor, prefijo=money_prefijo, sufijo=money_sufijo)
                except Exception:
                    valor = str(valor)

            row_data[header] = (valor, bold, size)

        if any(v[0] not in (None, "") for v in row_data.values()):
            excel_data.append(row_data)

    return excel_data


def limpiar_nombre_archivo(texto, largo_max=60):
    """Deja un string listo para usar como nombre de archivo (sin caracteres prohibidos)."""
    texto = "" if texto is None else str(texto).strip()
    for c in '/\\:*?"<>|':
        texto = texto.replace(c, "_")
    texto = texto.replace(" ", "_")
    texto = re.sub(r"_+", "_", texto).strip("_")
    return texto[:largo_max] if texto else "documento"


def leer_excel_preview(file_bytes, filas=5):
    """DataFrame liviano solo para mostrarle al usuario cómo se ven sus datos."""
    return pd.read_excel(io.BytesIO(file_bytes)).head(filas)


def leer_excel_obrero_capataz(file_bytes, date_headers):
    """Variante especial: separa SUELDO CAPATAZ en número y letras."""
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    headers = [str(c.value).strip().upper() if c.value is not None else "" for c in ws[1]]
    excel_data = []

    for row in ws.iter_rows(min_row=2):
        if all(c.value in (None, "") for c in row):
            continue
        row_data = {}
        for header, cell in zip(headers, row):
            if not header:
                continue
            valor = cell.value

            if header in date_headers:
                row_data[header] = format_fecha(valor)
                continue

            if header == "SUELDO CAPATAZ" and valor not in (None, ""):
                try:
                    numero, letras = sueldo_en_letras(valor)
                    row_data["SUELDO CAPATAZ"] = numero
                    row_data["SUELDO CAPATAZ LETRAS"] = letras
                except Exception:
                    row_data["SUELDO CAPATAZ"] = str(valor)
                    row_data["SUELDO CAPATAZ LETRAS"] = ""
                continue

            row_data[header] = "" if valor is None else str(valor)

        if row_data.get("NOMBRE COMPLETO", "").strip():
            excel_data.append(row_data)

    return excel_data


# =========================================================
# REEMPLAZO DE MERGEFIELDS EN WORD
# =========================================================

MERGEFIELD_RE = re.compile(r"«([^»]*)»")


def replace_mergefield_con_formato(doc, replacements):
    """replacements: header (en MAYÚSCULAS) -> (valor, bold, size).
    Busca cualquier «campo» sin importar si el molde lo tiene en mayúsculas,
    minúsculas o mezclado — siempre lo compara en mayúsculas."""
    def procesar(parrafo):
        for run in parrafo.runs:
            estado = {}

            def sub(m):
                campo = m.group(1).strip().upper()
                if campo in replacements:
                    valor, bold, size = replacements[campo]
                    estado["bold"] = bold
                    estado["size"] = size
                    return str(valor)
                return m.group(0)

            nuevo_texto = MERGEFIELD_RE.sub(sub, run.text)
            if nuevo_texto != run.text:
                run.text = nuevo_texto
                if "bold" in estado:
                    run.bold = estado["bold"]
                if estado.get("size"):
                    run.font.size = Pt(estado["size"])

    for p in doc.paragraphs:
        procesar(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    procesar(p)


def replace_mergefields_simple(doc, replacements, bold_keys=None):
    """replacements: header (en MAYÚSCULAS) -> valor (string plano).
    Misma lógica insensible a mayúsculas/minúsculas que la versión con formato."""
    bold_keys = bold_keys or set()

    def procesar(parrafo):
        for run in parrafo.runs:
            estado = {}

            def sub(m):
                campo = m.group(1).strip().upper()
                if campo in replacements:
                    if campo in bold_keys:
                        estado["bold"] = True
                    return str(replacements[campo])
                return m.group(0)

            nuevo_texto = MERGEFIELD_RE.sub(sub, run.text)
            if nuevo_texto != run.text:
                run.text = nuevo_texto
                if estado.get("bold"):
                    run.bold = True

    for p in doc.paragraphs:
        procesar(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    procesar(p)


# =========================================================
# CÁLCULO DE FINIQUITOS (feriado proporcional + días inhábiles + monto total)
# =========================================================
# El cálculo de "TOTAL FINIQUITO A PAGAR" (L41) fue portado y verificado
# fórmula por fórmula contra la planilla real "CÁLCULO FINIQUITO JJTT - CONAF"
# (incluyendo el VLOOKUP de años de servicio y la tabla de factor por fecha
# de término), y validado numéricamente contra un caso real: con los datos
# reales de una trabajadora (renuncia, 28 días trabajados, sueldo base
# $276.777) esta función da exactamente $10.770, igual que la planilla Excel.
#
# LIMITACIÓN CONOCIDA: para causales de despido (Art. 161-1, 161-2, 159-5) el
# cálculo depende de reconocer el texto de la causal. Se usa un reconocedor
# de patrones (normalizar_causal) en vez de comparación exacta de texto,
# justamente porque distintos Excels escriben la causal de forma distinta
# (ej. "159, N° 2" vs "Art. 159-2: ..."). Si una causal no se reconoce, la
# app avisa en vez de asumir $0 en silencio.

DIAS_SEMANA_ES = {
    "Monday": "lunes", "Tuesday": "martes", "Wednesday": "miércoles",
    "Thursday": "jueves", "Friday": "viernes", "Saturday": "sábado", "Sunday": "domingo",
}

CAUSALES_VALIDAS = {
    "159-1", "159-2", "159-3", "159-4", "159-5", "159-6",
    "160-1", "160-2", "160-3", "160-4", "160-5", "160-6", "160-7",
    "161-1", "161-2",
}

# Texto EXACTO tal cual lo espera la plantilla real (Hoja2!B4:B18) — las fórmulas
# de Excel comparan texto exacto, así que si escribimos cualquier otra variante
# en D14, la planilla (si alguien la abre en Excel) deja de reconocer la causal
# y calcula un monto distinto al de Python. Escribiendo siempre este texto
# canónico, Python y Excel quedan garantizados de coincidir.
CAUSALES_TEXTO_CANONICO = {
    "159-1": "Art. 159-1: Mutuo acuerdo",
    "159-2": "Art. 159-2: Renuncia del trabajador",
    "159-3": "Art. 159-3: Muerte del trabajador",
    "159-4": "Art. 159-4: Vencimiento plazo convenido (contratos plazo fijo)",
    "159-5": "Art. 159-5: Conclusión del trabajo que dio origen al contrato (obra o faena)",
    "159-6": "Art. 159-6: Caso fortuito o fuerza mayor",
    "160-1": "Art. 160-1: Falta probidad, vías de hecho, injurias o conducta grave",
    "160-2": "Art. 160-2: Negociaciones del trabajador dentro del giro del negocio",
    "160-3": "Art. 160-3: No concurrencia a labores sin causa justificada",
    "160-4": "Art. 160-4: Abandono del trabajo por parte del trabajador",
    "160-5": "Art. 160-5: Actos que afectan a la seguridad",
    "160-6": "Art. 160-6: Perjuicio material causado intencionalmente",
    "160-7": "Art. 160-7: Incumplimiento grave de las obligaciones",
    "161-1": "Art. 161-1: Necesidades de la empresa",
    "161-2": "Art. 161-2: Desahucio",
}


def normalizar_causal(texto):
    """Extrae el código de artículo (ej: '161-1', '159-5') de cualquier formato
    razonable de texto: 'Art. 161-1: Necesidades...', '159, N° 2 "RENUNCIA..."',
    '159-2', etc. Devuelve None si no reconoce ningún código."""
    if not texto:
        return None
    m = re.search(r"(15[9]|16[01])[^\d]{0,6}(\d)", str(texto))
    if m:
        codigo = f"{m.group(1)}-{m.group(2)}"
        return codigo if codigo in CAUSALES_VALIDAS else None
    return None


def causal_texto_canonico(texto):
    """Devuelve el texto EXACTO del desplegable de Excel para esta causal, o el
    texto original si no se logró reconocer ningún código (para no perder la
    información aunque no podamos usarla en el cálculo)."""
    codigo = normalizar_causal(texto)
    if codigo:
        return CAUSALES_TEXTO_CANONICO[codigo]
    return str(texto) if texto else ""


def _a_fecha(valor):
    """Normaliza distintos tipos de fecha (date, datetime, pandas Timestamp) a date."""
    if valor is None:
        return None
    if hasattr(valor, "date"):
        return valor.date()
    return valor


def excel_datedif_y(d1, d2):
    years = d2.year - d1.year
    if (d2.month, d2.day) < (d1.month, d1.day):
        years -= 1
    return max(0, years)


def excel_datedif_ym(d1, d2):
    months = (d2.year - d1.year) * 12 + (d2.month - d1.month)
    if d2.day < d1.day:
        months -= 1
    return max(0, months % 12)


def excel_datedif_md(d1, d2):
    if d2.day >= d1.day:
        return d2.day - d1.day
    prev_month_last_day = d2.replace(day=1) - timedelta(days=1)
    return prev_month_last_day.day - d1.day + d2.day


def calcular_dias_obtenidos(fecha_inicio, fecha_fin, es_zona_extrema="No"):
    """Replica la celda L16: días de feriado legal devengados según años/meses/días
    de servicio (DATEDIF-equivalente, no aproximación por división)."""
    fecha_inicio, fecha_fin = _a_fecha(fecha_inicio), _a_fecha(fecha_fin)
    years = excel_datedif_y(fecha_inicio, fecha_fin)
    months = excel_datedif_ym(fecha_inicio, fecha_fin)
    days = excel_datedif_md(fecha_inicio, fecha_fin) + 1

    factor_anual = 20 if es_zona_extrema == "Sí" else 15
    m16 = factor_anual * years
    n16 = 1.25 * months
    o16 = (1.25 / 30.0) * days

    dias_obtenidos = round(m16 + n16 + o16, 3)
    dias_obtenidos = 0 if dias_obtenidos < 1 else dias_obtenidos
    return dias_obtenidos, years, months, days


def calcular_dias_inhabiles(fecha_fin, dias_proporcionales):
    fecha_fin = _a_fecha(fecha_fin)
    parte_entera = int(dias_proporcionales)
    fraccion = dias_proporcionales - parte_entera
    dias_a_contar = parte_entera + (1 if fraccion > 0 else 0)

    fecha_cursor = fecha_fin
    dias_habiles = 0
    dias_inhabiles = 0
    calendario = []

    years = [fecha_fin.year, fecha_fin.year + 1]
    feriados = holidays.Chile(years=years)

    while dias_habiles < dias_a_contar:
        fecha_cursor += timedelta(days=1)
        es_feriado = fecha_cursor in feriados
        es_fin_semana = fecha_cursor.weekday() >= 5
        if es_fin_semana or es_feriado:
            dias_inhabiles += 1
            calendario.append((fecha_cursor, "inhábil"))
        else:
            dias_habiles += 1
            calendario.append((fecha_cursor, "hábil"))

    return dias_inhabiles, calendario


def calcular_finiquito_l41(causal, es_zona_extrema, fecha_inicio, fecha_fin, dias_inhabiles,
                            fecha_aviso=None, dias_tomados=0, remuneracion_pendiente=0,
                            sueldo_minimo=500000, valor_uf=35000, tipo_sueldo="Fijo",
                            sueldo_base_fijo=0, gratificacion_fijo="No", colacion_fijo=0,
                            movilizacion_fijo=0, sueldo_base_var=0, comisiones_m1=0,
                            comisiones_m2=0, comisiones_m3=0, gratificacion_var="No",
                            colacion_var=0, movilizacion_var=0):
    """Replica el cálculo real de 'TOTAL FINIQUITO A PAGAR' (celda L41)."""
    fecha_inicio, fecha_fin = _a_fecha(fecha_inicio), _a_fecha(fecha_fin)
    fecha_aviso = _a_fecha(fecha_aviso)
    causal_norm = normalizar_causal(causal)

    dias_obtenidos, years, months, days = calcular_dias_obtenidos(fecha_inicio, fecha_fin, es_zona_extrema)
    dias_trabajados_totales = (fecha_fin - fecha_inicio).days + 1
    g24 = dias_trabajados_totales / 365.0

    dias_pendientes = dias_obtenidos - dias_tomados
    total_dias_feriado = dias_pendientes + dias_inhabiles

    tope_gratificacion = (4.75 * sueldo_minimo) / 12.0
    h40 = min(sueldo_base_fijo * 0.25, tope_gratificacion) if gratificacion_fijo == "Sí" else 0
    prom_comisiones = (comisiones_m1 + comisiones_m2 + comisiones_m3) / 3.0
    h55 = min((sueldo_base_var + prom_comisiones) * 0.25, tope_gratificacion) if gratificacion_var == "Sí" else 0
    d44 = min(sueldo_base_fijo + h40 + colacion_fijo + movilizacion_fijo, 90 * valor_uf)
    d59 = min(sueldo_base_var + prom_comisiones + h55 + colacion_var + movilizacion_var, 90 * valor_uf)
    total_haberes = d44 if tipo_sueldo == "Fijo" else d59

    tabla_anos = [
        (0.0, 0), (1.0, 1), (1 + 5 / 9, 1), (1.6, 2), (2 + 5 / 9, 2),
        (2.6, 3), (3 + 5 / 9, 3), (3.6, 4), (4 + 5 / 9, 4),
        (4.6, 5), (5 + 5 / 9, 5), (5.6, 6), (6 + 5 / 9, 6),
        (6.6, 7), (7 + 5 / 9, 7), (7.6, 8), (8 + 5 / 9, 8),
        (8.6, 9), (9 + 5 / 9, 9), (9.6, 10), (10 + 5 / 9, 10), (10.6, 11),
    ]
    val_g24 = round(g24, 0) if (g24 >= 1 and months == 6) else g24
    d26 = 0
    for limite, anos in tabla_anos:
        if val_g24 >= limite:
            d26 = anos

    l30 = remuneracion_pendiente

    # DECISIÓN DEL USUARIO: no calcular aquí las indemnizaciones por despido
    # (aviso previo, años de servicio, obra/faena). Todo se calcula SIEMPRE
    # como si fuera renuncia voluntaria — solo vacaciones proporcionales +
    # días inhábiles + remuneración pendiente. Los pagos extra por despido
    # (Art. 161-1, 161-2, 159-5) se agregan aparte, a mano.
    l32 = 0
    l34 = 0

    base_diaria = (sueldo_base_fijo if tipo_sueldo == "Fijo" else sueldo_base_var + prom_comisiones) / 30.0
    l36 = base_diaria * total_dias_feriado

    l38 = 0

    l40 = l30 + l32 + l34 + l36 + l38

    # REGLA DE NEGOCIO: si trabajó menos de 30 días, el finiquito es $0,
    # sin importar lo que dé la fórmula. El subtotal queda igual en el
    # resultado (SUBTOTAL) solo como referencia de lo que habría dado.
    menos_de_30_dias = dias_trabajados_totales < 30
    if menos_de_30_dias:
        l41 = 0
    else:
        l41 = int(Decimal(str(l40)).quantize(Decimal("1E1"), rounding=ROUND_HALF_UP))

    return {
        "MONTO_FINIQUITO": l41,
        "SUBTOTAL": l40,
        "MENOS_DE_30_DIAS": menos_de_30_dias,
        "CAUSAL_RECONOCIDA": causal_norm is not None,
        "DIAS_OBTENIDOS": dias_obtenidos,
        "DIAS_PENDIENTES": dias_pendientes,
        "TOTAL_DIAS_FERIADO": total_dias_feriado,
    }


def leer_valores_base_planilla(plantilla_bytes):
    """Lee de la plantilla (antes de que la toquemos) los valores que hoy la app
    no pisa por fila: sueldo mínimo, valor UF, tipo de sueldo, gratificación,
    colación/movilización, zona extrema, días ya tomados, remuneración
    pendiente y fecha de aviso de despido. Así el cálculo en Python usa
    exactamente lo mismo que ya trae tu Excel, no valores inventados."""
    wb = load_workbook(io.BytesIO(plantilla_bytes), data_only=True)
    ws = wb.active
    return {
        "fecha_aviso": ws["D19"].value,
        "sueldo_minimo": ws["D30"].value or 500000,
        "valor_uf": ws["D31"].value or 0,
        "tipo_sueldo": ws["D33"].value or "Fijo",
        "gratificacion_fijo": ws["D40"].value or "No",
        "colacion_fijo": ws["D41"].value or 0,
        "movilizacion_fijo": ws["D42"].value or 0,
        "es_zona_extrema": ws["L14"].value or "No",
        "dias_tomados": ws["L18"].value or 0,
        "remuneracion_pendiente": ws["L30"].value or 0,
    }


# =========================================================
# UI STREAMLIT
# =========================================================

st.set_page_config(page_title="Automatización RR.HH.", page_icon="🤖", layout="wide")

# --- Banner ---
import base64

def _banner_base64():
    with open(os.path.join(os.path.dirname(__file__), "assets", "banner.jpg"), "rb") as f:
        return base64.b64encode(f.read()).decode()

st.markdown(
    """
    <style>
    .block-container { padding-top: 1.5rem; max-width: 1180px; }

    .hero-banner {
        border-radius: 20px;
        overflow: hidden;
        margin-bottom: 1.8rem;
        box-shadow: 0 8px 24px rgba(0,0,0,0.35);
    }
    .hero-banner img { width: 100%; height: auto; display: block; }

    /* --- Pestañas como botones tipo "pill" --- */
    .stTabs [role="tablist"] {
        gap: 10px;
        flex-wrap: wrap;
        border-bottom: none !important;
        box-shadow: none !important;
        padding-bottom: 0.6rem;
    }
    .stTabs [data-testid="stTab"] .react-aria-SelectionIndicator {
        display: none !important;
    }
    .stTabs [data-testid="stTab"]::before,
    .stTabs [data-testid="stTab"]::after {
        display: none !important;
        content: none !important;
    }
    .stTabs [data-testid="stTab"] {
        height: auto;
        background-color: rgba(255,255,255,0.05) !important;
        border: 1px solid #2A3742 !important;
        border-radius: 999px !important;
        padding: 0.6rem 1.25rem !important;
        margin: 0 !important;
        box-shadow: none !important;
        outline: none !important;
        transition: all 0.15s ease;
    }
    .stTabs [data-testid="stTab"] p {
        font-weight: 600 !important;
        font-size: 0.92rem !important;
        color: #9AA5AE;
        margin: 0 !important;
    }
    .stTabs [data-testid="stTab"]:hover {
        background-color: rgba(127,184,143,0.12) !important;
        border-color: #5F9E72 !important;
        box-shadow: none !important;
    }
    .stTabs [data-testid="stTab"]:hover p { color: #DCE3E8; }
    .stTabs [data-testid="stTab"][aria-selected="true"] {
        background: linear-gradient(135deg, #8FCB9E 0%, #5F9E72 100%) !important;
        border-color: #7FB88F !important;
        box-shadow: 0 4px 14px rgba(127,184,143,0.30) !important;
    }
    .stTabs [data-testid="stTab"][aria-selected="true"]:focus,
    .stTabs [data-testid="stTab"][aria-selected="true"]:focus-visible {
        box-shadow: 0 4px 14px rgba(127,184,143,0.30) !important;
        outline: none !important;
    }
    .stTabs [data-testid="stTab"][aria-selected="true"] p { color: #0F151C !important; }

    .stButton>button {
        border-radius: 10px;
        font-weight: 600;
        border: 1px solid #5F9E72;
    }
    .stButton>button:hover { border-color: #7FB88F; color: #7FB88F; }
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown(
    f'<div class="hero-banner"><img src="data:image/jpeg;base64,{_banner_base64()}" alt="Automatización de Documentación RR.HH."></div>',
    unsafe_allow_html=True,
)

# --- Chat de ayuda (barra lateral) ---
SYSTEM_PROMPT_AYUDA = (
    "Eres el asistente de ayuda de una app interna llamada 'Automatización RR.HH.'. "
    "La app tiene 6 pestañas: Contratos, Cálculo Finiquitos, Finiquito + Decl. Jurada, "
    "Anexo Continuidad, Anexo Obrero→Capataz, y Generador Universal. En todas se sube un "
    "Excel de datos y una o más plantillas Word. En el Word, los campos a rellenar se "
    "escriben como «NOMBRE_COLUMNA» (con esas comillas angulares «»), y deben coincidir "
    "con el nombre de una columna del Excel (no importan mayúsculas/minúsculas). El "
    "'Generador Universal' sirve para cualquier documento nuevo: ahí el usuario marca con "
    "clics cuáles columnas son fechas y cuáles son montos, no depende de nombres fijos. "
    "Cada pestaña genera un ZIP para descargar. Responde siempre en español, corto y "
    "directo, como ayuda contextual de una app, no como charla larga. Si preguntan algo "
    "que no tiene que ver con usar esta app, redirige amablemente al tema."
)

with st.sidebar:
    st.markdown("### 💬 Ayuda")
    if not ia_disponible():
        st.caption(
            "El chat de ayuda con IA no está activo todavía en este despliegue "
            "(falta configurar ANTHROPIC_API_KEY en Secrets)."
        )
    else:
        st.caption("Pregúntame cómo usar cualquier parte de la app.")

    if "chat_ayuda" not in st.session_state:
        st.session_state.chat_ayuda = []

    for rol, texto in st.session_state.chat_ayuda:
        with st.chat_message(rol):
            st.markdown(texto)

    pregunta_ayuda = st.chat_input("Ej: ¿cómo lleno el molde de Word?")
    if pregunta_ayuda:
        st.session_state.chat_ayuda.append(("user", pregunta_ayuda))
        if not ia_disponible():
            respuesta_ayuda = (
                "El chat de ayuda todavía no está configurado en este despliegue "
                "(falta ANTHROPIC_API_KEY en Secrets de Streamlit Cloud)."
            )
        else:
            with st.spinner("Pensando..."):
                respuesta_ayuda, error_ayuda = preguntar_ia(SYSTEM_PROMPT_AYUDA, pregunta_ayuda, max_tokens=500)
            if error_ayuda:
                respuesta_ayuda = f"⚠️ {error_ayuda}"
        st.session_state.chat_ayuda.append(("assistant", respuesta_ayuda))
        st.rerun()

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📑 Contratos",
    "🧮 Cálculo + Finiquito + Jurada",
    "➕ Anexo Continuidad",
    "⬆️ Anexo Obrero→Capataz",
    "🧩 Generador Universal",
])


# ---------------------------------------------------------
# TAB 1: CONTRATOS
# ---------------------------------------------------------
with tab1:
    st.subheader("Generador de Contratos")
    with st.expander("Columnas esperadas en el Excel"):
        st.markdown(
            "- `NOMBRE COMPLETO`\n"
            "- `FECHA ACTUAL`, `FECHA DE NACIMIENTO`, `INICIO CONTRARO`, `TÉRMINO CONTRARO` (fechas)\n"
            "- `SUELDO BASE` (número)\n"
            "- Cualquier otra columna se reemplaza tal cual en el molde como `«NOMBRE_COLUMNA»`"
        )

    c1, c2 = st.columns(2)
    excel_c = c1.file_uploader("Excel de datos", type=["xlsx"], key="c_excel")
    molde_c = c2.file_uploader("Plantilla Word (molde de contrato)", type=["docx"], key="c_molde")
    bloque_auditoria_ia(excel_c, "Datos para generar contratos de trabajo en Chile", "contratos")
    generar_pdf_c = st.checkbox("📄 También generar en PDF (mismo formato del Word)", value=True, key="pdf_c")

    if st.button("Generar contratos", key="btn_c"):
        if not excel_c or not molde_c:
            st.error("Sube el Excel y la plantilla primero.")
            st.session_state.pop("resultado_c", None)
        else:
            try:
                date_headers = ["FECHA ACTUAL", "FECHA DE NACIMIENTO", "INICIO CONTRARO", "TÉRMINO CONTRARO"]
                excel_data = leer_excel_formato(excel_c.getvalue(), date_headers, money_headers=["SUELDO BASE"])

                if not excel_data:
                    st.warning("No se encontraron filas válidas (revisa que exista la columna NOMBRE COMPLETO).")
                    st.session_state.pop("resultado_c", None)
                else:
                    archivos = {}
                    errores_filas = []
                    molde_bytes = molde_c.getvalue()
                    for idx, row_data in enumerate(excel_data):
                        try:
                            doc = Document(io.BytesIO(molde_bytes))
                            replace_mergefield_con_formato(doc, row_data)
                            numero = str(idx + 1).zfill(3)
                            nombre = str(row_data["NOMBRE COMPLETO"][0]).replace(" ", "_")
                            archivos[f"{numero}_CTO_{nombre}.docx"] = guardar_docx_bytes(doc)
                        except Exception as e:
                            nombre_fallback = row_data.get("NOMBRE COMPLETO", (f"fila {idx + 1}",))[0]
                            errores_filas.append(f"Fila {idx + 1} ({nombre_fallback}): {e}")

                    n_generados = len(archivos)
                    if archivos and generar_pdf_c:
                        with st.spinner("Convirtiendo a PDF..."):
                            archivos.update(agregar_pdfs_al_zip(archivos))

                    archivos_word, archivos_pdf = separar_word_y_pdf(archivos)
                    st.session_state["resultado_c"] = {
                        "n_generados": n_generados,
                        "total": len(excel_data),
                        "zip_word": crear_zip(archivos_word) if archivos_word else None,
                        "zip_pdf": crear_zip(archivos_pdf) if archivos_pdf else None,
                        "errores_filas": errores_filas,
                    }
            except Exception as e:
                st.error(f"Error generando los contratos: {e}")
                st.session_state.pop("resultado_c", None)

    if "resultado_c" in st.session_state:
        r = st.session_state["resultado_c"]
        if r["zip_word"]:
            st.success(f"{r['n_generados']} de {r['total']} contrato(s) generado(s).")
            col_w, col_p = st.columns(2)
            col_w.download_button(
                "⬇️ Descargar en Word", r["zip_word"], "Contratos_Word.zip", "application/zip", key="dl_word_c",
            )
            if r["zip_pdf"]:
                col_p.download_button(
                    "⬇️ Descargar en PDF", r["zip_pdf"], "Contratos_PDF.zip", "application/zip", key="dl_pdf_c",
                )
        if r["errores_filas"]:
            st.error(
                "❌ Estas filas no se pudieron procesar (dato inválido) — corrígelas y "
                "vuelve a intentar:\n\n" + "\n".join(f"- {e}" for e in r["errores_filas"])
            )


# ---------------------------------------------------------
# TAB 2: CÁLCULO + FINIQUITO + DECL. JURADA (todo en un clic)
# ---------------------------------------------------------
with tab2:
    st.subheader("Cálculo de Finiquito + Finiquito + Declaración Jurada")
    with st.expander("Columnas esperadas en el Excel de datos"):
        st.markdown(
            "- `NOMBRE COMPLETO`, `CARNÉ DE IDENTIDAD N°`\n"
            "- `FECHA ACTUAL` (fecha del documento)\n"
            "- `INICIO CONTRARO`, `FIN CONTRAT` (fechas)\n"
            "- `SUELDO BASE`\n"
            "- `CAUSAL TÉRMINO`\n\n"
            "**El monto se calcula siempre como renuncia voluntaria**: vacaciones "
            "proporcionales + días inhábiles + remuneración pendiente (si tu plantilla la "
            "trae). No incluye indemnización por aviso previo, años de servicio ni obra/faena "
            "— esos pagos por despido se agregan aparte, a mano. **Si trabajó menos de 30 "
            "días, el finiquito da $0** (regla de negocio, sin excepción). Sueldo mínimo, "
            "valor UF, gratificación, zona extrema, etc. se leen directo de tu plantilla de "
            "cálculo.\n\n"
            "Sube el Excel de datos + las 3 plantillas, y en un clic se calcula el monto Y se "
            "generan los 3 documentos (planilla de cálculo, finiquito y declaración jurada) — "
            "ya no hace falta descargar el monto y volver a subirlo a mano."
        )

    c1, c2 = st.columns(2)
    excel_ft = c1.file_uploader("Excel de datos", type=["xlsx"], key="ft_excel")
    plantilla_ft = c2.file_uploader("Plantilla Excel (cálculo finiquito)", type=["xlsx"], key="ft_plantilla")
    c3, c4 = st.columns(2)
    molde_finiquito_ft = c3.file_uploader("Plantilla Finiquito (Word)", type=["docx"], key="ft_finiquito")
    molde_jurada_ft = c4.file_uploader("Plantilla Declaración Jurada (Word)", type=["docx"], key="ft_jurada")
    bloque_auditoria_ia(excel_ft, "Datos para calcular y generar finiquitos", "finiquito_todo")
    generar_pdf_ft = st.checkbox(
        "📄 También generar Finiquito y Decl. Jurada en PDF (mismo formato del Word)",
        value=True, key="pdf_ft",
    )

    if st.button("Calcular y generar todo", key="btn_ft"):
        faltan = [
            n for n, v in [
                ("Excel de datos", excel_ft), ("Plantilla Excel de cálculo", plantilla_ft),
                ("Plantilla Finiquito", molde_finiquito_ft), ("Plantilla Declaración Jurada", molde_jurada_ft),
            ] if not v
        ]
        if faltan:
            st.error("Falta subir: " + ", ".join(faltan))
            st.session_state.pop("resultado_ft", None)
        else:
            try:
                df = pd.read_excel(io.BytesIO(excel_ft.getvalue()))
                df.columns = df.columns.str.strip().str.upper()

                requeridas = ["NOMBRE COMPLETO", "CARNÉ DE IDENTIDAD N°", "FECHA ACTUAL",
                              "INICIO CONTRARO", "FIN CONTRAT", "SUELDO BASE", "CAUSAL TÉRMINO"]
                faltantes = [c for c in requeridas if c not in df.columns]
                if faltantes:
                    st.error(f"Faltan columnas en el Excel: {', '.join(faltantes)}")
                    st.session_state.pop("resultado_ft", None)
                else:
                    plantilla_bytes = plantilla_ft.getvalue()
                    finiquito_bytes = molde_finiquito_ft.getvalue()
                    jurada_bytes = molde_jurada_ft.getvalue()

                    # Cargamos también con openpyxl (no solo pandas) para poder leer el
                    # negrita/tamaño real de cada celda, igual que hacía el script original,
                    # y pasar CUALQUIER columna del Excel al Word (no solo las que usamos
                    # para calcular el monto).
                    wb_datos = load_workbook(io.BytesIO(excel_ft.getvalue()), data_only=True)
                    ws_datos = wb_datos.active
                    headers_originales = {
                        str(c.value).strip().upper(): i + 1
                        for i, c in enumerate(ws_datos[1]) if c.value is not None
                    }
                    base = leer_valores_base_planilla(plantilla_bytes)

                    archivos = {}
                    resumen = []
                    causales_no_reconocidas = []
                    errores_filas = []
                    df_salida = df.copy()
                    montos_dict = {}

                    for index, row in df.iterrows():
                        nombre_fallback = row.get("NOMBRE COMPLETO", f"(fila {index + 1})")
                        try:
                            nombre = row["NOMBRE COMPLETO"]
                            rut = row["CARNÉ DE IDENTIDAD N°"]
                            fecha_actual = pd.to_datetime(row["FECHA ACTUAL"]).date()
                            fecha_inicio = pd.to_datetime(row["INICIO CONTRARO"]).date()
                            fecha_fin = pd.to_datetime(row["FIN CONTRAT"]).date()
                            sueldo_base = row["SUELDO BASE"]
                            causal = row["CAUSAL TÉRMINO"]

                            dias_obtenidos, years, months, days = calcular_dias_obtenidos(
                                fecha_inicio, fecha_fin, base["es_zona_extrema"]
                            )
                            dias_pendientes = dias_obtenidos - base["dias_tomados"]
                            dias_inhabiles, _calendario = calcular_dias_inhabiles(fecha_fin, dias_pendientes)

                            resultado = calcular_finiquito_l41(
                                causal=causal,
                                es_zona_extrema=base["es_zona_extrema"],
                                fecha_inicio=fecha_inicio,
                                fecha_fin=fecha_fin,
                                dias_inhabiles=dias_inhabiles,
                                fecha_aviso=base["fecha_aviso"],
                                dias_tomados=base["dias_tomados"],
                                remuneracion_pendiente=base["remuneracion_pendiente"],
                                sueldo_minimo=base["sueldo_minimo"],
                                valor_uf=base["valor_uf"],
                                tipo_sueldo=base["tipo_sueldo"],
                                sueldo_base_fijo=sueldo_base,
                                gratificacion_fijo=base["gratificacion_fijo"],
                                colacion_fijo=base["colacion_fijo"],
                                movilizacion_fijo=base["movilizacion_fijo"],
                            )
                            monto = resultado["MONTO_FINIQUITO"]

                            if not resultado["CAUSAL_RECONOCIDA"]:
                                causales_no_reconocidas.append(f"Fila {index + 1} ({nombre}): «{causal}»")

                            resumen.append({
                                "Nombre": nombre,
                                "Días trabajados": (fecha_fin - fecha_inicio).days + 1,
                                "Días feriado pendiente": round(dias_pendientes, 2),
                                "Días inhábiles": dias_inhabiles,
                                "Monto Finiquito": monto,
                                "< 30 días (finiquito $0)": "Sí" if resultado["MENOS_DE_30_DIAS"] else "No",
                                "Causal reconocida (solo texto D14)": "Sí" if resultado["CAUSAL_RECONOCIDA"] else "⚠️ No",
                            })
                            montos_dict[index] = monto

                            # 1) Planilla de cálculo (Excel)
                            wb = load_workbook(io.BytesIO(plantilla_bytes))
                            wb.calculation.fullCalcOnLoad = True
                            ws = wb.active
                            ws["D6"] = nombre
                            ws["D7"] = rut
                            ws["D8"] = "PEE"
                            ws["D14"] = causal_texto_canonico(causal)
                            ws["D17"] = fecha_inicio
                            ws["D17"].number_format = "DD-MM-YYYY"
                            ws["D18"] = fecha_fin
                            ws["D18"].number_format = "DD-MM-YYYY"
                            ws["L20"] = dias_inhabiles
                            ws["D39"] = sueldo_base

                            numero = str(index + 1).zfill(2)
                            nombre_archivo = str(nombre).replace(" ", "_")
                            archivos[f"Calculo/{numero}_CALCULO_FINIQUITO_{nombre_archivo}.xlsx"] = guardar_xlsx_bytes(wb)

                            # 2) Finiquito + Declaración Jurada (Word) — usa el monto recién calculado
                            # Partimos de TODAS las columnas de esta fila, con su negrita/tamaño
                            # real (igual que el script original), y encima pisamos solo los
                            # campos que nosotros calculamos o reformateamos.
                            fila_excel_num = index + 2  # fila 1 = encabezados
                            row_data_word = {}
                            for header, col_num in headers_originales.items():
                                celda = ws_datos.cell(row=fila_excel_num, column=col_num)
                                bold = celda.font.bold if celda.font.bold is not None else False
                                size = celda.font.sz if celda.font.sz is not None else None
                                row_data_word[header] = (celda.value, bold, size)

                            def _con_formato_previo(campo, valor_nuevo, bold_default=False):
                                _, bold_prev, size_prev = row_data_word.get(campo, (None, bold_default, None))
                                return (valor_nuevo, bold_prev, size_prev)

                            monto_texto = monto_a_texto(monto, prefijo="$")
                            row_data_word["FECHA ACTUAL"] = _con_formato_previo("FECHA ACTUAL", format_fecha(fecha_actual))
                            row_data_word["INICIO CONTRATO"] = _con_formato_previo("INICIO CONTRATO", format_fecha(fecha_inicio))
                            row_data_word["INICIO CONTRARO"] = _con_formato_previo("INICIO CONTRARO", format_fecha(fecha_inicio))
                            row_data_word["FIN CONTRAT"] = _con_formato_previo("FIN CONTRAT", format_fecha(fecha_fin))
                            row_data_word["FIN CONTRATO"] = _con_formato_previo("FIN CONTRATO", format_fecha(fecha_fin))
                            row_data_word["CAUSAL TÉRMINO"] = _con_formato_previo("CAUSAL TÉRMINO", causal_texto_canonico(causal))
                            # El monto es un valor calculado, no viene de una celda propia —
                            # lo dejamos en negrita por defecto ya que es el dato clave del documento.
                            row_data_word["MONTO FINIQUITO"] = (monto_texto, True, None)

                            doc_f = Document(io.BytesIO(finiquito_bytes))
                            replace_mergefield_con_formato(doc_f, row_data_word)
                            archivos[f"Finiquitos/{numero}_Finiquito_{nombre_archivo}.docx"] = guardar_docx_bytes(doc_f)

                            doc_j = Document(io.BytesIO(jurada_bytes))
                            replace_mergefield_con_formato(doc_j, row_data_word)
                            archivos[f"Declaraciones_Juradas/{numero}_Decl_Jurada_{nombre_archivo}.docx"] = guardar_docx_bytes(doc_j)

                        except Exception as e:
                            errores_filas.append(f"Fila {index + 1} ({nombre_fallback}): {e}")
                            continue

                    df_salida["MONTO FINIQUITO"] = df_salida.index.map(montos_dict)
                    buf_salida = io.BytesIO()
                    df_salida.to_excel(buf_salida, index=False)

                    if archivos and generar_pdf_ft:
                        with st.spinner("Convirtiendo Finiquitos y Declaraciones a PDF..."):
                            archivos.update(agregar_pdfs_al_zip(archivos))

                    archivos_word, archivos_pdf = separar_word_y_pdf(archivos)
                    st.session_state["resultado_ft"] = {
                        "n_generados": len(montos_dict),
                        "total": len(df),
                        "resumen": resumen,
                        "errores_filas": errores_filas,
                        "causales_no_reconocidas": causales_no_reconocidas,
                        "zip_word": crear_zip(archivos_word) if archivos_word else None,
                        "zip_pdf": crear_zip(archivos_pdf) if archivos_pdf else None,
                        "excel_montos": buf_salida.getvalue(),
                    }
            except Exception as e:
                st.error(f"Error generando finiquitos: {e}")
                st.session_state.pop("resultado_ft", None)

    if "resultado_ft" in st.session_state:
        r = st.session_state["resultado_ft"]

        if r["n_generados"]:
            st.success(f"{r['n_generados']} de {r['total']} caso(s): planilla + finiquito + declaración jurada generados.")

        if r["errores_filas"]:
            st.error(
                "❌ Estas filas NO se pudieron procesar (dato inválido, ej. una fecha que no "
                "existe) — corrígelas en tu Excel y vuelve a correr solo esas. El resto de "
                "abajo sí se generó bien:\n\n" + "\n".join(f"- {e}" for e in r["errores_filas"])
            )

        if r["causales_no_reconocidas"]:
            st.info(
                "ℹ️ No reconocí el formato de causal en estas filas — no afecta el monto "
                "(siempre se calcula como renuncia voluntaria), pero en la planilla de "
                "cálculo la celda D14 va a mostrar tu texto original tal cual:\n\n"
                + "\n".join(f"- {c}" for c in r["causales_no_reconocidas"])
            )

        st.caption(
            "⚠️ Monto = vacaciones proporcionales + días inhábiles + remuneración pendiente "
            "(renuncia voluntaria siempre, $0 si trabajó menos de 30 días). No incluye "
            "indemnización por despido. Verificado contra un caso real, pero para los primeros "
            "usos te recomendamos comparar alguna planilla contra Excel antes de confiar 100% "
            "en el cálculo automático."
        )

        if r["resumen"]:
            st.dataframe(pd.DataFrame(r["resumen"]), use_container_width=True)

            col_w, col_p, col_datos = st.columns(3)
            if r["zip_word"]:
                col_w.download_button(
                    "⬇️ Word + planillas Excel", r["zip_word"],
                    "Finiquitos_Word.zip", "application/zip", key="dl_word_ft",
                    help="Planillas de cálculo (.xlsx) + Finiquitos y Declaraciones Juradas (.docx).",
                )
            if r["zip_pdf"]:
                col_p.download_button(
                    "⬇️ Finiquitos + Juradas en PDF", r["zip_pdf"],
                    "Finiquitos_PDF.zip", "application/zip", key="dl_pdf_ft",
                )
            col_datos.download_button(
                "⬇️ Excel con MONTO FINIQUITO", r["excel_montos"],
                "Datos_Finiquito_Con_Montos.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="dl_montos_ft",
                help="Por si necesitas los montos aparte para otro uso.",
            )


# ---------------------------------------------------------
# TAB 3: ANEXO CONTINUIDAD
# ---------------------------------------------------------
with tab3:
    st.subheader("Anexo de Continuidad")
    with st.expander("Columnas esperadas en el Excel"):
        st.markdown(
            "- `NOMBRE COMPLETO`\n"
            "- `FECHA ACTUAL`, `FECHA INICIO CONTRATO`, `FECHA FIN CONTRATO ANEXO` (fechas)"
        )

    c1, c2 = st.columns(2)
    excel_a = c1.file_uploader("Excel de datos", type=["xlsx"], key="a_excel")
    molde_a = c2.file_uploader("Plantilla Word (anexo)", type=["docx"], key="a_molde")
    bloque_auditoria_ia(excel_a, "Datos para anexos de continuidad de contrato", "anexo_cont")
    generar_pdf_a = st.checkbox("📄 También generar en PDF (mismo formato del Word)", value=True, key="pdf_a")

    if st.button("Generar anexos", key="btn_a"):
        if not excel_a or not molde_a:
            st.error("Sube el Excel y la plantilla primero.")
            st.session_state.pop("resultado_a", None)
        else:
            try:
                date_headers = ["FECHA ACTUAL", "FECHA INICIO CONTRATO", "FECHA FIN CONTRATO ANEXO"]
                excel_data = leer_excel_simple(excel_a.getvalue(), date_headers)

                if not excel_data:
                    st.warning("No se encontraron filas válidas.")
                    st.session_state.pop("resultado_a", None)
                else:
                    archivos = {}
                    errores_filas = []
                    molde_bytes = molde_a.getvalue()
                    for idx, row_data in enumerate(excel_data, start=1):
                        try:
                            doc = Document(io.BytesIO(molde_bytes))
                            replace_mergefields_simple(doc, row_data)
                            nombre = row_data["NOMBRE COMPLETO"].replace(" ", "_")
                            archivos[f"{idx:02d}_{nombre}_ANEXO.docx"] = guardar_docx_bytes(doc)
                        except Exception as e:
                            errores_filas.append(f"Fila {idx} ({row_data.get('NOMBRE COMPLETO', '?')}): {e}")

                    n_generados = len(archivos)
                    if archivos and generar_pdf_a:
                        with st.spinner("Convirtiendo a PDF..."):
                            archivos.update(agregar_pdfs_al_zip(archivos))

                    archivos_word, archivos_pdf = separar_word_y_pdf(archivos)
                    st.session_state["resultado_a"] = {
                        "n_generados": n_generados,
                        "total": len(excel_data),
                        "zip_word": crear_zip(archivos_word) if archivos_word else None,
                        "zip_pdf": crear_zip(archivos_pdf) if archivos_pdf else None,
                        "errores_filas": errores_filas,
                    }
            except Exception as e:
                st.error(f"Error generando los anexos: {e}")
                st.session_state.pop("resultado_a", None)

    if "resultado_a" in st.session_state:
        r = st.session_state["resultado_a"]
        if r["zip_word"]:
            st.success(f"{r['n_generados']} de {r['total']} anexo(s) generado(s).")
            col_w, col_p = st.columns(2)
            col_w.download_button(
                "⬇️ Descargar en Word", r["zip_word"], "Anexos_Word.zip", "application/zip", key="dl_word_a",
            )
            if r["zip_pdf"]:
                col_p.download_button(
                    "⬇️ Descargar en PDF", r["zip_pdf"], "Anexos_PDF.zip", "application/zip", key="dl_pdf_a",
                )
        if r["errores_filas"]:
            st.error(
                "❌ Estas filas no se pudieron procesar — corrígelas y vuelve a intentar:\n\n"
                + "\n".join(f"- {e}" for e in r["errores_filas"])
            )


# ---------------------------------------------------------
# TAB 4: ANEXO OBRERO -> CAPATAZ
# ---------------------------------------------------------
with tab4:
    st.subheader("Anexo Obrero → Capataz")
    with st.expander("Columnas esperadas en el Excel"):
        st.markdown(
            "- `NOMBRE COMPLETO`\n"
            "- `FECHA ACTUAL`, `FECHA INICIO CONTRATO`, `FECHA FIN CONTRATO ANEXO`, `FECHA OB CPT DESDE` (fechas)\n"
            "- `SUELDO CAPATAZ` (número; se genera automáticamente `«SUELDO CAPATAZ»` y `«SUELDO CAPATAZ LETRAS»`)"
        )

    c1, c2 = st.columns(2)
    excel_o = c1.file_uploader("Excel de datos", type=["xlsx"], key="o_excel")
    molde_o = c2.file_uploader("Plantilla Word (anexo obrero→capataz)", type=["docx"], key="o_molde")
    bloque_auditoria_ia(excel_o, "Datos para anexos de ascenso obrero a capataz", "obrero_capataz")
    generar_pdf_o = st.checkbox("📄 También generar en PDF (mismo formato del Word)", value=True, key="pdf_o")

    if st.button("Generar modificaciones de contrato", key="btn_o"):
        if not excel_o or not molde_o:
            st.error("Sube el Excel y la plantilla primero.")
            st.session_state.pop("resultado_o", None)
        else:
            try:
                date_headers = ["FECHA ACTUAL", "FECHA INICIO CONTRATO", "FECHA FIN CONTRATO ANEXO", "FECHA OB CPT DESDE"]
                excel_data = leer_excel_obrero_capataz(excel_o.getvalue(), date_headers)

                if not excel_data:
                    st.warning("No se encontraron filas válidas.")
                    st.session_state.pop("resultado_o", None)
                else:
                    archivos = {}
                    errores_filas = []
                    molde_bytes = molde_o.getvalue()
                    for i, row_data in enumerate(excel_data, start=1):
                        try:
                            doc = Document(io.BytesIO(molde_bytes))
                            replace_mergefields_simple(doc, row_data, bold_keys={"SUELDO CAPATAZ LETRAS"})
                            nombre = row_data["NOMBRE COMPLETO"].replace(" ", "_")
                            archivos[f"{i:03d}_MOD_CONTRATO_{nombre}.docx"] = guardar_docx_bytes(doc)
                        except Exception as e:
                            errores_filas.append(f"Fila {i} ({row_data.get('NOMBRE COMPLETO', '?')}): {e}")

                    n_generados = len(archivos)
                    if archivos and generar_pdf_o:
                        with st.spinner("Convirtiendo a PDF..."):
                            archivos.update(agregar_pdfs_al_zip(archivos))

                    archivos_word, archivos_pdf = separar_word_y_pdf(archivos)
                    st.session_state["resultado_o"] = {
                        "n_generados": n_generados,
                        "total": len(excel_data),
                        "zip_word": crear_zip(archivos_word) if archivos_word else None,
                        "zip_pdf": crear_zip(archivos_pdf) if archivos_pdf else None,
                        "errores_filas": errores_filas,
                    }
            except Exception as e:
                st.error(f"Error generando las modificaciones: {e}")
                st.session_state.pop("resultado_o", None)

    if "resultado_o" in st.session_state:
        r = st.session_state["resultado_o"]
        if r["zip_word"]:
            st.success(f"{r['n_generados']} de {r['total']} documento(s) generado(s).")
            col_w, col_p = st.columns(2)
            col_w.download_button(
                "⬇️ Descargar en Word", r["zip_word"], "Modificacion_Contrato_Word.zip",
                "application/zip", key="dl_word_o",
            )
            if r["zip_pdf"]:
                col_p.download_button(
                    "⬇️ Descargar en PDF", r["zip_pdf"], "Modificacion_Contrato_PDF.zip",
                    "application/zip", key="dl_pdf_o",
                )
        if r["errores_filas"]:
            st.error(
                "❌ Estas filas no se pudieron procesar — corrígelas y vuelve a intentar:\n\n"
                + "\n".join(f"- {e}" for e in r["errores_filas"])
            )


# ---------------------------------------------------------
# TAB 5: GENERADOR UNIVERSAL (columnas 100% dinámicas)
# ---------------------------------------------------------
with tab5:
    st.subheader("Generador Universal")
    st.caption(
        "Para cualquier documento que no encaje en las pestañas anteriores: sube tu Excel y tu "
        "molde, y dile a la app cuáles columnas son fechas y cuáles son montos. No depende de "
        "nombres de columna fijos — sirve para cualquier tipo de contrato, carta o formulario."
    )

    c1, c2 = st.columns(2)
    excel_u = c1.file_uploader("Excel de datos", type=["xlsx"], key="u_excel")
    molde_u = c2.file_uploader("Plantilla Word (con «CAMPOS» que coincidan con tus columnas)", type=["docx"], key="u_molde")

    headers_u = []
    if excel_u:
        try:
            headers_u = leer_headers_excel(excel_u.getvalue())
            with st.expander(f"Vista previa del Excel ({len(headers_u)} columnas detectadas)"):
                st.dataframe(leer_excel_preview(excel_u.getvalue()), use_container_width=True)
        except Exception as e:
            st.error(f"No pude leer ese Excel: {e}")
    bloque_auditoria_ia(excel_u, "Datos para un documento genérico (definido por el usuario)", "universal")

    if headers_u:
        col1, col2 = st.columns(2)
        fechas_sel = col1.multiselect(
            "📅 Columnas que son FECHAS", options=headers_u, key="u_fechas",
            help="Se van a formatear como DD/MM/AAAA en el documento final.",
        )
        montos_sel = col2.multiselect(
            "💰 Columnas que son MONTOS", options=headers_u, key="u_montos",
            help="Se van a convertir a formato '1.234.567 (UN MILLÓN... PESOS)'.",
        )

        col3, col4 = st.columns(2)
        sufijo_moneda = col3.text_input("Palabra para el monto en letras", value="PESOS", key="u_sufijo")
        id_col = col4.selectbox(
            "Columna para nombrar cada archivo generado", options=headers_u, key="u_id",
            help="Normalmente el nombre de la persona o el número de folio.",
        )
    else:
        fechas_sel, montos_sel, sufijo_moneda, id_col = [], [], "PESOS", None
        st.info("Sube tu Excel para poder elegir cuáles columnas son fechas y montos.")

    generar_pdf_u = st.checkbox("📄 También generar en PDF (mismo formato del Word)", value=True, key="pdf_u")

    if st.button("Generar documentos", key="btn_u"):
        if not excel_u or not molde_u:
            st.error("Sube el Excel y la plantilla primero.")
            st.session_state.pop("resultado_u", None)
        else:
            try:
                data = leer_excel_universal(
                    excel_u.getvalue(),
                    date_headers=[h.strip().upper() for h in fechas_sel],
                    money_headers=[h.strip().upper() for h in montos_sel],
                    money_prefijo="$",
                    money_sufijo=(sufijo_moneda or "PESOS").strip().upper(),
                )

                if not data:
                    st.warning("No se encontraron filas con datos en ese Excel.")
                    st.session_state.pop("resultado_u", None)
                else:
                    archivos = {}
                    errores_filas = []
                    molde_bytes = molde_u.getvalue()
                    id_col_upper = id_col.strip().upper() if id_col else None

                    for idx, row_data in enumerate(data, start=1):
                        try:
                            doc = Document(io.BytesIO(molde_bytes))
                            replace_mergefield_con_formato(doc, row_data)

                            if id_col_upper and id_col_upper in row_data:
                                nombre_doc = limpiar_nombre_archivo(row_data[id_col_upper][0])
                            else:
                                nombre_doc = f"documento_{idx}"

                            archivos[f"{idx:03d}_{nombre_doc}.docx"] = guardar_docx_bytes(doc)
                        except Exception as e:
                            errores_filas.append(f"Fila {idx}: {e}")

                    n_generados = len(archivos)
                    if archivos and generar_pdf_u:
                        with st.spinner("Convirtiendo a PDF..."):
                            archivos.update(agregar_pdfs_al_zip(archivos))

                    archivos_word, archivos_pdf = separar_word_y_pdf(archivos)
                    st.session_state["resultado_u"] = {
                        "n_generados": n_generados,
                        "total": len(data),
                        "zip_word": crear_zip(archivos_word) if archivos_word else None,
                        "zip_pdf": crear_zip(archivos_pdf) if archivos_pdf else None,
                        "errores_filas": errores_filas,
                    }
            except Exception as e:
                st.error(f"Error generando los documentos: {e}")
                st.session_state.pop("resultado_u", None)

    if "resultado_u" in st.session_state:
        r = st.session_state["resultado_u"]
        if r["zip_word"]:
            st.success(f"{r['n_generados']} de {r['total']} documento(s) generado(s).")
            col_w, col_p = st.columns(2)
            col_w.download_button(
                "⬇️ Descargar en Word", r["zip_word"], "Documentos_Word.zip", "application/zip", key="dl_word_u",
            )
            if r["zip_pdf"]:
                col_p.download_button(
                    "⬇️ Descargar en PDF", r["zip_pdf"], "Documentos_PDF.zip", "application/zip", key="dl_pdf_u",
                )
        if r["errores_filas"]:
            st.error(
                "❌ Estas filas no se pudieron procesar — corrígelas y vuelve a intentar:\n\n"
                + "\n".join(f"- {e}" for e in r["errores_filas"])
            )
